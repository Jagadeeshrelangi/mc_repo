import os
import docx
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

# Define base paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
KNOWLEDGE_BASE_DIR = os.path.join(BASE_DIR, "knowledge_base")
INDEX_DIR = os.path.join(KNOWLEDGE_BASE_DIR, "faiss_index")

# Ensure required folders exist
folders = ["cars", "bikes", "manuals", "maintenance", "dashboard_symbols", "repair_guides", "obd_codes", "faq"]
for f in folders:
    os.makedirs(os.path.join(KNOWLEDGE_BASE_DIR, f), exist_ok=True)

def load_text_or_md(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def load_pdf(path: str) -> str:
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def load_docx(path: str) -> str:
    doc = docx.Document(path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def build_vector_index():
    print("--- Scanning Knowledge Base for Documents ---")
    documents = []
    
    # Supported file loaders map
    loaders = {
        ".txt": load_text_or_md,
        ".md": load_text_or_md,
        ".pdf": load_pdf,
        ".docx": load_docx
    }
    
    # Recursively traverse knowledge base directory
    for root, _, files in os.walk(KNOWLEDGE_BASE_DIR):
        category = os.path.basename(root)
        if category == "faiss_index":
            continue
            
        for file in files:
            file_path = os.path.join(root, file)
            _, ext = os.path.splitext(file)
            ext = ext.lower()
            
            if ext in loaders:
                try:
                    print(f"Loading [{category}] {file} ...")
                    content = loaders[ext](file_path)
                    
                    if content.strip():
                        # Create LangChain Document with metadata context
                        doc = Document(
                            page_content=content,
                            metadata={"source": file, "category": category}
                        )
                        documents.append(doc)
                except Exception as e:
                    print(f"Failed to load {file}: {str(e)}")
                    
    if not documents:
        print("No documents found to index.")
        return
        
    print(f"\nLoaded {len(documents)} source documents.")
    
    # Chunking
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    split_docs = text_splitter.split_documents(documents)
    print(f"Split documents into {len(split_docs)} chunks.")
    
    # Embeddings Initialization (Local CPU model)
    print("Initializing embedding model (sentence-transformers/all-MiniLM-L6-v2) ...")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    
    # FAISS compilation
    print("Building FAISS index ...")
    db = FAISS.from_documents(split_docs, embeddings)
    
    # Save index files
    os.makedirs(INDEX_DIR, exist_ok=True)
    db.save_local(INDEX_DIR)
    print(f"FAISS index successfully saved to: {INDEX_DIR}")

if __name__ == "__main__":
    build_vector_index()
